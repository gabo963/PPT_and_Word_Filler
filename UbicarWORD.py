import win32com.client
from PIL import ImageGrab
import os
import sys
import shutil
import datetime
import docx
import openpyxl

rutaFotos = ''

def diccionario( book ):
    sheet = book['Encontrar Parámetros Word']
    
    fila = 3
    columna = 2
    
    i = 0
    tamano = 0
    
    while sheet.cell( row=(i+fila), column=(columna) ).value != None:
        i += 1
    
    tamano = i
    
    dic = {}
    fots = {}
    
    for i in range( tamano ):
        dic[sheet.cell( row=(i+fila), column=(columna) ).value] = sheet.cell( row=(i+fila), column=(columna+2) ).value
        fots[sheet.cell( row=(i+fila), column=(columna) ).value] = sheet.cell( row=(i+fila), column=(columna+1) ).value
        
    return [dic, fots]
    
def guardaGraficosComoImagenes(inputExcelFilePath, hoja):

    global rutaFotos

    rutaFotos = inputExcelFilePath[:inputExcelFilePath.rindex('/')]+ '/Fotos/'

    if os.path.exists( rutaFotos ):
        outputPNGImagePath = rutaFotos
    else:
        os.makedirs( rutaFotos )
        outputPNGImagePath = rutaFotos
    
    mydir = win32com.__gen_path__
    
    if os.path.exists(mydir):
        shutil.rmtree(mydir)

    o = win32com.client.Dispatch("Excel.Application")

    wb = o.Workbooks.Open(inputExcelFilePath)

    sheet = o.Sheets(hoja)
    for n, shape in enumerate(sheet.ChartObjects()):
        
        
        shape.Copy()
        image = ImageGrab.grabclipboard()
        
        image.save(str(outputPNGImagePath+shape.Name+'.png'), 'png')
        pass
    pass

    wb.Close(True)
    o.Quit()


def ubicarPalabraW( rutaWord, rutaExcel, me ):
    
    documento = docx.Document( rutaWord )
    book = openpyxl.load_workbook( filename=rutaExcel, data_only=True )
    
    params = diccionario( book )
    
    guardaGraficosComoImagenes(rutaExcel, 'Encontrar Parámetros Word')    

    global rutaFotos

    dic = params[0]
    fots = params[1]    
    
    for line in documento.paragraphs:
        nuevo = line
        for val in dic:
            if val in line.text:
                
                if fots[val] == "Foto":
                        
                        line.text = line.text.replace( val, "" )
                        nuevo.runs[-1].add_break()
                        nuevo.runs[-1].add_picture( rutaFotos + dic[val] + '.png' )
                else:
                    if isinstance(dic[val], datetime.date):
                        line.text = line.text.replace( val, str(dic[val].strftime('%d/%m/%y')) )
                    else:
                        line.text = line.text.replace( val, str(dic[val]) )
            
    documento.save( rutaWord[:len(rutaWord)-5] + ' ' +list(dic.values())[0] + '.docx')
    
    me['text'] = 'Hecho'
    me['state'] = 'disabled'
    me['command'] = None
    
             
